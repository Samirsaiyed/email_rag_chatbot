"""
Create sample attachments in multiple formats (PDF, DOCX, TXT).
"""
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from pathlib import Path
from src.config import ATTACHMENTS_DIR

def create_pdf(filename: str, content: list):
    """Create a PDF with given content."""
    pdf_path = ATTACHMENTS_DIR / "pdfs" / filename
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter
    
    for page_num, page_content in enumerate(content, 1):
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, f"Page {page_num}")
        
        c.setFont("Helvetica", 11)
        y_position = height - 100
        
        for line in page_content.split('\n'):
            if y_position < 50:
                break
            c.drawString(50, y_position, line)
            y_position -= 15
        
        c.showPage()
    
    c.save()
    print(f"Created PDF: {filename}")

def create_docx(filename: str, title: str, content: str):
    """Create a DOCX file."""
    docx_path = ATTACHMENTS_DIR / "docx" / filename
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    doc.add_heading(title, 0)
    
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    doc.save(str(docx_path))
    print(f"Created DOCX: {filename}")

def create_txt(filename: str, content: str):
    """Create a TXT file."""
    txt_path = ATTACHMENTS_DIR / "txt" / filename
    txt_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"Created TXT: {filename}")

def main():
    """Create attachments for multiple threads."""
    
    print("Creating sample attachments (PDF, DOCX, TXT)...\n")
    
    # ===== Thread T-58ae003b: Storage/Budget (existing 5 PDFs + add 1 DOCX) =====
    create_docx(
        "Storage_Upgrade_Notes.docx",
        "Storage Upgrade - Internal Notes",
        """Meeting Notes: Storage Vendor Selection
Date: April 5, 2001
Attendees: John Doe, Sarah Johnson, Mike Chen

Key Discussion Points:

We reviewed three vendor proposals for the storage upgrade project. The consensus is leaning toward StorageTech Solutions based on their superior technical specifications and competitive pricing.

Budget Considerations:
The initial budget request was $38,000, but after reviewing the technical requirements, we've increased the recommendation to $45,000. This includes installation, migration services, and extended warranty.

Timeline:
- Vendor selection: April 15
- Contract signing: April 30
- Installation: May 15-30
- Migration: June 1-15
- Go-live: June 20

Action Items:
- John: Finalize contract negotiations
- Sarah: Process purchase requisition
- Mike: Prepare migration plan

Next Steps:
Need executive approval for the $45,000 budget. Once approved, we can proceed with contract execution."""
    )
    
    # ===== Thread T-3df8a268: Axia Energy (2 PDFs + 1 DOCX + 1 TXT) =====
    create_pdf("Axia_Energy_Partnership_Agreement.pdf", [
        """Axia Energy Partnership Agreement
Date: February 15, 2001

Parties:
- Enron Corporation
- Axia Energy, LP

Terms:
1. Natural Gas Supply: 50,000 MMBtu/day
2. Contract Duration: 36 months
3. Pricing: NYMEX + $0.15/MMBtu
4. Delivery Point: Houston Ship Channel

Financial Terms:
- Monthly payment: Estimated $2.5M
- Security deposit: $500,000
- Payment terms: Net 30 days

Approved by: Legal Department
Signed: February 20, 2001"""
    ])
    
    create_pdf("Axia_Credit_Analysis.pdf", [
        """Credit Analysis Report
Subject: Axia Energy, LP
Date: February 8, 2001
Analyst: Risk Management Team

Credit Rating: BBB+
Credit Limit Recommendation: $5,000,000

Financial Highlights:
- Annual Revenue: $180M
- Assets: $95M
- Debt-to-Equity: 1.2
- Current Ratio: 1.8

Risk Factors:
- Concentration in Texas market
- Limited hedging program
- Seasonal demand variability

Recommendation: Approve with standard credit monitoring""",
        
        """Page 2

Historical Performance:
- Payment history: Excellent (no late payments)
- Years in business: 8
- Management experience: Strong

Collateral:
- Gas storage facilities: $12M
- Pipeline capacity rights: $8M

Monitoring Requirements:
- Quarterly financial statements
- Monthly usage reports
- Annual credit review"""
    ])
    
    create_docx(
        "Axia_Negotiation_Notes.docx",
        "Axia Energy - Negotiation Summary",
        """Negotiation Summary: Axia Energy Partnership
Date: February 12, 2001
Lead Negotiator: Commercial Trading Desk

Overview:
Successful negotiation with Axia Energy for long-term natural gas supply agreement. Final terms represent a win-win arrangement that secures reliable supply while maintaining acceptable pricing.

Key Negotiation Points:

Volume Commitment:
Initial request was 75,000 MMBtu/day. Settled at 50,000 MMBtu/day with option to increase to 65,000 after 12 months based on performance.

Pricing Structure:
Axia wanted fixed pricing. We pushed for index-based pricing. Final agreement: NYMEX + $0.15/MMBtu basis differential, which is favorable given current market conditions.

Contract Duration:
Compromised on 36 months (their ask: 60 months, our initial offer: 24 months). This provides adequate stability while maintaining flexibility.

Credit Terms:
Agreed to $500,000 security deposit (down from their initial ask of $1M). Credit limit approved at $5M based on risk analysis.

Delivery Points:
Secured Houston Ship Channel as primary delivery point, which aligns with our distribution infrastructure.

Next Steps:
- Legal review: February 13-14
- Executive approval: February 15
- Contract execution: February 20
- First delivery: March 1"""
    )
    
    create_txt(
        "Axia_Email_Forward.txt",
        """From: Trading Desk
To: Risk Management
Date: February 1, 2001
Subject: FW: Axia Energy - Initial Proposal

---------- Forwarded message ----------
From: Axia Energy <contact@axiaenergy.com>
Date: January 29, 2001
Subject: Partnership Proposal

Dear Enron Team,

Axia Energy LP is interested in establishing a long-term natural gas supply partnership with Enron. We are a growing energy company based in Houston with proven reserves and strong operational capabilities.

Key highlights of our proposal:

- Supply volume: Up to 75,000 MMBtu/day
- Duration: 5-year agreement preferred
- Delivery: Houston Ship Channel area
- Pricing: Competitive fixed rates or index-based (open to discussion)

We believe a partnership would be mutually beneficial given:
1. Our reliable supply capabilities
2. Enron's distribution network
3. Aligned geographic focus in Texas market

Would appreciate the opportunity to discuss this proposal in detail. Available for meeting next week.

Best regards,
David Martinez
VP Business Development
Axia Energy, LP
"""
    )
    
    # ===== Thread T-8b62a250: El Paso Electric (2 PDFs + 1 TXT) =====
    create_pdf("El_Paso_Power_Purchase_Agreement.pdf", [
        """Power Purchase Agreement
Seller: Enron Power Marketing
Buyer: El Paso Electric Company
Date: March 20, 2001

Power Delivery Terms:
- Capacity: 150 MW
- Delivery Period: June 2001 - May 2002
- Delivery Point: Palo Verde hub

Pricing:
- Base Rate: $45/MWh
- Peak Hours (6am-10pm): $55/MWh
- Off-Peak: $35/MWh

Total Contract Value: Approximately $48M annually

Transmission:
- Buyer responsible for transmission costs
- Seller arranges scheduling

Force Majeure: Standard clauses apply
Termination: 90 days notice required"""
    ])
    
    create_pdf("El_Paso_Load_Forecast.pdf", [
        """Load Forecast Analysis
Customer: El Paso Electric Company
Period: Summer 2001
Date: March 5, 2001

Peak Demand Forecast:
- June: 1,250 MW
- July: 1,450 MW
- August: 1,480 MW
- September: 1,320 MW

Weather Assumptions:
- Average temperatures 2°F above normal
- 15% probability of extreme heat event

Load Growth:
- Year-over-year: 3.5%
- New commercial: 25 MW
- Residential growth: 15 MW""",
        
        """Page 2

Reserve Margin Analysis:
- Current capacity: 1,650 MW
- Required reserves: 15%
- Shortfall risk: Low

Renewable Integration:
- Solar: 45 MW (planned)
- Wind: 30 MW (existing)

Recommendations:
- Secure additional 100 MW for peak season
- Consider interruptible contracts
- Monitor weather forecasts closely"""
    ])
    
    create_txt(
        "El_Paso_Pricing_Discussion.txt",
        """INTERNAL MEMO - El Paso Electric Pricing Strategy
From: West Power Trading Desk
Date: March 18, 2001

Pricing Discussion Summary:

Customer requested fixed pricing at $42/MWh for entire contract. We countered with time-of-day pricing to better reflect market conditions and manage our risk.

Final Agreed Pricing:
- Peak (6am-10pm): $55/MWh
- Off-peak: $35/MWh  
- Blended average: ~$45/MWh

Rationale:
1. Summer peak demand in Southwest drives high spot prices
2. Our generation portfolio benefits from peak pricing
3. Customer gets predictability while we maintain margin flexibility

Comparison to Market:
- Spot market average (summer): $60-80/MWh peak
- Our pricing represents 20-30% discount to spot
- Competitive vs. other suppliers (verified through market intel)

Risk Mitigation:
- Hedged 80% of expected delivery through forward purchases
- Remaining 20% exposure acceptable given portfolio diversity
- Weather derivatives considered for extreme heat scenarios

Contract Value: ~$48M annually
Expected Margin: 12-15%
"""
    )
    
    # ===== Thread T-a5f23567: PG&E (2 PDFs + 1 DOCX) =====
    create_pdf("PGE_California_Crisis_Memo.pdf", [
        """Internal Memo - California Energy Crisis
To: Trading Desk
From: West Power Desk
Date: January 25, 2001
RE: PG&E Exposure and Market Conditions

URGENT: PG&E Credit Situation

Current Exposure:
- Outstanding receivables: $185M
- Mark-to-market exposure: $67M
- Total at-risk: $252M

Credit Actions Taken:
- Reduced credit line from $500M to $100M
- Required daily collateral posting
- Halted new forward transactions

California Market Conditions:
- Spot prices: $200-$400/MWh (normal: $30-50)
- Reserve margin: Critical (<5%)
- Rolling blackouts: Stage 2 alerts frequent"""
    ])
    
    create_pdf("PGE_Risk_Mitigation_Strategy.pdf", [
        """Risk Mitigation Strategy
Subject: PG&E Bankruptcy Risk
Date: January 30, 2001

Immediate Actions:
1. Cease new physical delivery contracts
2. Liquidate forward positions where possible
3. Increase collateral requirements to 120%
4. Daily exposure reporting to senior management

Scenario Analysis:
- Best case: State bailout, exposure recovered
- Base case: Partial recovery 60-70%
- Worst case: Chapter 11, recovery 30-40%

Hedging Recommendations:
- Purchase credit default protection
- Diversify California counterparties
- Reduce overall CA market exposure by 40%""",
        
        """Page 2

Legal Considerations:
- Netting agreements: Review enforceability
- Preference period: 90 days pre-filing
- Setoff rights: Confirm jurisdiction

Financial Impact:
- Potential loss: $75M - $175M
- Impact on Q1 earnings: Significant
- Reserve requirements: $100M recommended

Next Steps:
- Daily credit committee briefings
- Coordinate with legal team
- Prepare disclosure for SEC filing"""
    ])
    
    create_docx(
        "PGE_Crisis_Timeline.docx",
        "PG&E Crisis - Key Events Timeline",
        """California Energy Crisis: PG&E Timeline
Compiled: February 5, 2001

December 2000:
- Dec 7: First rolling blackouts in California
- Dec 14: California ISO declares Stage 3 emergency
- Dec 15: PG&E announces severe financial stress

January 2001:
- Jan 2: Spot prices exceed $300/MWh
- Jan 10: PG&E credit rating downgraded to junk
- Jan 16: Governor declares state of emergency
- Jan 17: State begins buying power for utilities
- Jan 25: Our credit committee reduces PG&E exposure limit

Key Issues:

Wholesale Prices:
Market prices have increased 10x from 2000 levels due to supply constraints, transmission bottlenecks, and strong demand.

Retail Rate Freeze:
AB 1890 prevents utilities from passing costs to consumers, creating massive losses for PG&E and SCE.

Credit Exposure:
Our receivables have grown to $185M. Bankruptcy risk is material and increasing daily.

State Intervention:
State purchasing power on behalf of utilities. Long-term solution unclear. Political pressure for rate increases vs. public backlash.

Our Strategy:
1. Minimize new California exposure
2. Collect receivables aggressively  
3. Prepare for potential bankruptcy filing
4. Monitor legislative developments

Outlook:
Situation remains fluid. Bankruptcy filing by PG&E within 60-90 days is probable unless state provides financial support."""
    )
    
    print(f"\n" + "="*60)
    print("SUMMARY:")
    print("="*60)
    print("\nThread T-58ae003b (Storage/Budget):")
    print("  - 5 PDFs (existing)")
    print("  - 1 DOCX (meeting notes)")
    print("\nThread T-3df8a268 (Axia Energy):")
    print("  - 2 PDFs (agreement, credit analysis)")
    print("  - 1 DOCX (negotiation notes)")
    print("  - 1 TXT (email forward)")
    print("\nThread T-8b62a250 (El Paso Electric):")
    print("  - 2 PDFs (PPA, load forecast)")
    print("  - 1 TXT (pricing discussion)")
    print("\nThread T-a5f23567 (PG&E):")
    print("  - 2 PDFs (crisis memo, risk strategy)")
    print("  - 1 DOCX (timeline)")
    print("\nTOTAL: 9 PDFs + 3 DOCX + 2 TXT = 14 attachments across 4 threads")

if __name__ == "__main__":
    main()